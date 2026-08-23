from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document


class DOCXLoader:
    """Load a DOCX file into a LangChain Document."""

    def load(
        self,
        file_path: str | Path,
        source_id: str,
        file_name: str | None = None,
    ) -> list[Document]:

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {path}")

        if path.suffix.lower() != ".docx":
            raise ValueError("DOCXLoader requires a .docx file")

        doc = DocxDocument(str(path))

        paragraphs = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        text = "\n\n".join(paragraphs)

        if not text:
            return []

        display_name = file_name or path.name

        return [
            Document(
                page_content=text,
                metadata={
                    "source_id": source_id,
                    "source_type": "docx",
                    "file_name": display_name,
                },
            )
        ]